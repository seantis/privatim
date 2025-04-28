from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from babel.dates import format_datetime
from privatim.i18n import translate, _
from privatim.layouts.layout import DEFAULT_TIMEZONE
from privatim.models.association_tables import AttendanceStatus
from privatim.utils import datetime_format
from pyramid.renderers import render
import lxml.html
import html2text
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from weasyprint import HTML, CSS  # type: ignore
from weasyprint.text.fonts import FontConfiguration  # type: ignore


from typing import TYPE_CHECKING, Protocol
if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph
    from privatim.models.association_tables import MeetingUserAttendance
    from privatim.models import Meeting
    from pytz import BaseTzInfo
    from pyramid.interfaces import IRequest


@dataclass
class ReportOptions:

    language: str = 'de'
    """ Translate report into language. """

    tz: 'BaseTzInfo' = DEFAULT_TIMEZONE
    """ Use timezone for all timestamps. """


@dataclass
class PDFDocument:

    data: bytes
    """ Raw PDF content. """

    filename: str
    """ Document's filename when stored or downloaded. """


class ReportRenderer(Protocol):

    def render(
        self, meeting: 'Meeting', timestamp: str, request: 'IRequest'
    ) -> bytes: ...


class MeetingReport:

    def __init__(
        self,
        request: 'IRequest',
        meeting: 'Meeting',
        options: ReportOptions,
        renderer: ReportRenderer,
    ) -> None:
        self.request = request
        self.meeting = meeting
        self.renderer = renderer
        self.options = options
        # Use the renderer's timestamp if available, otherwise generate one
        self.issued_at = getattr(renderer, 'issued_at', datetime.utcnow())

    @property
    def created_at(self) -> str:
        """Returns formatted report date."""
        return format_datetime(
            self.issued_at,
            format='short',
            locale=self.options.language,
            tzinfo=self.options.tz,
        ).replace('\u202f', ' ')

    @property
    def filename(self) -> str:
        datestr = format_datetime(
            self.issued_at,
            format='YYMMdd',
            locale=self.options.language,
            tzinfo=self.options.tz,
        )
        # Filename extension depends on the renderer used
        extension = getattr(self.renderer, 'extension', 'pdf')
        return f'{datestr}_{self.meeting.name}.{extension}'

    def build(self) -> PDFDocument:
        """Render report using the provided renderer."""

        pdf = self.renderer.render(self.meeting, self.created_at, self.request)

        return PDFDocument(pdf, self.filename)


class HTMLReportRenderer:
    """
    Render meeting report with WeasyPrint, using HTML and CSS.
    You can turn on logging for weasyprint to debug issues:
        >>> import logging, sys
        >>> logger = logging.getLogger('weasyprint')
        >>> logger.setLevel(logging.DEBUG)
        >>> logger.addHandler(logging.StreamHandler(sys.stdout))
    """

    template = 'privatim:reporting/template/report.pt'

    def render(
        self, meeting: 'Meeting', timestamp: str, request: 'IRequest'
    ) -> bytes:
        html = self.render_template(meeting, timestamp, request)
        return self.render_pdf(html, request)

    def render_template(
        self, meeting: 'Meeting', timestamp: str, request: 'IRequest'
    ) -> str:
        """Render chameleon report template."""
        document_context = {'title': meeting.name, 'created_at': timestamp}
        title = translate(
            _(
                "Protocol of meeting ${title}",
                mapping={'title': document_context['title']},
            )
        )
        ctx = {
            'title': title,
            'meeting': meeting,
            'sorted_attendance_records': meeting.attendance_records,
            'meeting_time': datetime_format(meeting.time),
            'document': document_context,
        }
        return render(self.template, ctx)

    def render_pdf(self, html: str, request: 'IRequest') -> bytes:
        """
        Render processed chameleon template as PDF.
        """
        resource_base_url = Path.cwd() / 'privatim' / 'reporting'
        buffer = BytesIO()

        italic_font_name = 'dm-sans-v6-latin-ext_latin-500italic.woff'
        italic_font_name_woff2 = 'dm-sans-v6-latin-ext_latin-500italic.woff2'
        regular_font_name = 'dm-sans-v6-latin-ext_latin-500.woff'
        regular_font_name_woff2 = 'dm-sans-v6-latin-ext_latin-500.woff2'
        normal_font_name = 'dm-sans-v6-latin-ext_latin-regular.woff'
        normal_font_name_woff2 = 'dm-sans-v6-latin-ext_latin-regular.woff2'
        normal_italic_font_name = 'dm-sans-v6-latin-ext_latin-italic.woff'
        normal_italic_font_name_wo2 = 'dm-sans-v6-latin-ext_latin-italic.woff2'
        base_font_url = 'privatim:static/fonts/'

        # Create URLs first
        font_paths = {
            'regular': f'{base_font_url}{regular_font_name}',
            'regular_woff2': f'{base_font_url}{regular_font_name_woff2}',
            'italic': f'{base_font_url}{italic_font_name}',
            'italic_woff2': f'{base_font_url}{italic_font_name_woff2}',
            'normal': f'{base_font_url}{normal_font_name}',
            'normal_woff2': f'{base_font_url}{normal_font_name_woff2}',
            'normal_italic': f'{base_font_url}{normal_italic_font_name}',
            'normal_italic_woff2': f'{base_font_url}'
            f'{normal_italic_font_name_wo2}',
        }

        font_urls = {
            key: request.static_url(path) for key, path in font_paths.items()
        }

        font_config = FontConfiguration()
        css_font_face = f'''
        @font-face {{
            font-family: 'DM Sans';
            font-style: normal;
            font-weight: 400;
            src: url({font_urls['normal_woff2']}) format('woff2'),
                 url({font_urls['normal']}) format('woff');
        }}
        @font-face {{
            font-family: 'DM Sans';
            font-style: italic;
            font-weight: 400;
            src: url({font_urls['normal_italic_woff2']}) format('woff2'),
                 url({font_urls['normal_italic']}) format('woff');
        }}
        @font-face {{
            font-family: 'DM Sans';
            font-style: normal;
            font-weight: 500;
            src: url({font_urls['regular_woff2']}) format('woff2'),
                 url({font_urls['regular']}) format('woff');
        }}
        @font-face {{
            font-family: 'DM Sans';
            font-style: italic;
            font-weight: 500;
            src: url({font_urls['italic_woff2']}) format('woff2'),
                 url({font_urls['italic']}) format('woff');
        }}
        '''

        css = CSS(string=css_font_face, font_config=font_config)
        HTML(string=html, base_url=str(resource_base_url)).write_pdf(
            buffer, stylesheets=[css], font_config=font_config
        )
        return buffer.getvalue()


def add_markdown_runs(paragraph: 'Paragraph', markdown_text: str) -> None:
    """
    Parses a simple markdown string (bold/italic) and adds formatted runs
    to the given python-docx paragraph.
    Handles **bold**, __bold__, *italic*, _italic_.
    Does not handle nesting or complex markdown.
    """
    # Split by markdown markers, keeping the markers as delimiters
    # Regex handles **, *, __, _
    parts = re.split(r'(\*\*|\*|__|_(?![_]))', markdown_text)

    is_bold = False
    is_italic = False

    for part in parts:
        if not part:  # Skip empty strings from split
            continue

        if part == '**' or part == '__':
            is_bold = not is_bold
        elif part == '*' or part == '_':
            is_italic = not is_italic
        else:  # Actual text content
            run = paragraph.add_run(part)
            run.bold = is_bold
            run.italic = is_italic


class WordReportRenderer:
    """
    Render meeting report as a Word (.docx) document using python-docx.
    Converts HTML descriptions to Markdown and then renders basic formatting.
    """

    extension = 'docx'

    def __init__(self) -> None:
        self.issued_at = datetime.utcnow()  # Store creation time

    def render(
        self, meeting: 'Meeting', timestamp: str, request: 'IRequest'
    ) -> bytes:
        """Generates the DOCX file content."""
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'DM Sans'  # Match PDF font if possible
        font.size = Pt(10)

        title_text = translate(
            _(
                "Protocol of meeting ${title}",
                mapping={'title': meeting.name},
            ),
            language=request.locale_name,
        )
        p = document.add_paragraph()
        p.add_run(title_text).bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_after = Pt(18)

        document.add_paragraph(
            f"{translate(_('Date / Time:'), language=request.locale_name)} "
            f"{datetime_format(meeting.time)}"
        )
        document.add_paragraph(
            f"{translate(_('Working Group:'), language=request.locale_name)} "
            f"{meeting.working_group.name}"
        )
        document.add_paragraph()  # Spacer

        # Attendees
        if meeting.attendance_records:
            document.add_paragraph(
                f"{translate(_('Attendees:'), language=request.locale_name)}"
            ).runs[0].bold = True
            # Fetch sorted records using the property
            sorted_records: list['MeetingUserAttendance'] = (
                request.dbsession.execute(meeting.sorted_attendance_records)
                .scalars()
                .all()
            )

            for record in sorted_records:
                status_marker = (
                    " ✓" if record.status == AttendanceStatus.ATTENDED else ""
                )
                document.add_paragraph(
                    f"{record.user.fullname}{status_marker}",
                    style='List Bullet',
                )
            document.add_paragraph() 

        document.add_paragraph(
            translate(_('Agenda Items'), language=request.locale_name)
        ).runs[0].bold = True
        document.add_paragraph()  # Spacer

        for index, item in enumerate(meeting.agenda_items, start=1):
            p_title = document.add_paragraph()
            p_title.add_run(f"{index}. ").bold = True
            p_title.add_run(item.title).bold = True
            p_title.paragraph_format.space_after = Pt(6)

            # Agenda Item Description (convert HTML->Markdown->DOCX runs)
            if item.description:
                last_paragraph = None
                try:
                    h = html2text.HTML2Text()
                    h.body_width = 0  # Disable line wrapping
                    h.ignore_links = True  # Ignore links for simplicity
                    h.ignore_images = True  # Ignore images
                    # Add other config as needed, e.g., h.ignore_emphasis = False
                    markdown_text = h.handle(item.description).strip()

                    # Split into paragraphs based on double newlines
                    md_paragraphs = markdown_text.split('\n\n')

                    for md_paragraph_text in md_paragraphs:
                        md_paragraph_text = md_paragraph_text.strip()
                        if not md_paragraph_text:
                            continue

                        # Add a new paragraph for each block from markdown
                        p_desc = document.add_paragraph()
                        p_desc.paragraph_format.left_indent = Inches(0.25)
                        # Spacing between paragraphs within a description
                        p_desc.paragraph_format.space_after = Pt(6)
                        add_markdown_runs(p_desc, md_paragraph_text)
                        last_paragraph = p_desc

                    # Ensure space after the entire description block
                    if last_paragraph:
                        last_paragraph.paragraph_format.space_after = Pt(12)

                except Exception:
                    # Fallback to plain text extraction using lxml
                    try:
                        html_tree = lxml.html.fromstring(item.description)
                        text_content = html_tree.text_content()
                        p_desc = document.add_paragraph(text_content.strip())
                        p_desc.paragraph_format.left_indent = Inches(0.25)
                        p_desc.paragraph_format.space_after = Pt(12)
                    except lxml.etree.ParserError:
                        p_desc = document.add_paragraph(
                            item.description
                        )  # Raw
                        p_desc.paragraph_format.left_indent = Inches(0.25)
                        p_desc.paragraph_format.space_after = Pt(12)

        # Save to buffer
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
